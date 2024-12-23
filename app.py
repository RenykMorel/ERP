from flask import Flask, session, render_template, jsonify, request, redirect, url_for, flash, render_template_string, current_app, send_file
from flask_login import LoginManager, current_user, login_user, logout_user, login_required, UserMixin
from facturas.facturas_models import Facturacion, PreFactura, NotaCredito, NotaDebito, Cliente
from inventario.inventario_models import InventarioItem, MovimientoInventario, ItemFactura
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_wtf.csrf import CSRFProtect, CSRFError
from config import Config
import requests
import json
from dotenv import load_dotenv
import sys
import os
from datetime import datetime, date
from sqlalchemy.exc import IntegrityError
from werkzeug.security import generate_password_hash, check_password_hash
import psycopg2
from psycopg2 import pool
import threading
import queue
from flask_migrate import Migrate
sys.path.append(os.path.join(os.path.dirname(__file__), 'banco'))
from banco.banco_models import NuevoBanco as Banco
from models import (Usuario, Empresa, Rol, Permiso, Modulo, UsuarioModulo, 
                    Notificacion, Transaccion, Cuenta, AdminPlanSuscripcion, 
                    AdminFactura, AdminConfiguracionSeguridad, AdminReporte)
from models import Transaccion, Cuenta
from mailjet_rest import Client
import secrets
import string
from logging.config import dictConfig
from flask import Blueprint
from extensions import db, login_manager, limiter, migrate
from blinker import signal
from sqlalchemy import event
import logging
from logging.handlers import RotatingFileHandler
from sqlalchemy import inspect
from sqlalchemy.orm import joinedload
from typing import Any, Optional
from typing import Tuple, Dict, Any, List
from flask import send_from_directory
from models import Usuario
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from io import BytesIO, StringIO
from flask import send_file, make_response, jsonify
import xlsxwriter
import base64
import csv
from reportlab.lib import colors
from common.models import ItemFactura, ItemPreFactura, MovimientoInventario
from asistentes.banco_assistant import AsistenteBancario
from asistentes.inventario_assistant import AsistenteInventario
from decimal import Decimal



# Importar el nuevo módulo de marketing
from marketing.routes import marketing

# Importar la biblioteca de OpenAI
from openai import OpenAI

load_dotenv()

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

mailjet = Client(auth=(os.getenv('MJ_APIKEY_PUBLIC'), os.getenv('MJ_APIKEY_PRIVATE')), version='v3.1')

def send_welcome_email(email, nombre, apellido, login_link, nombre_usuario, password):
    nombre_completo = f"{nombre} {apellido}"
    
    html_content = render_template_string("""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Bienvenido a CalculAI</title>
        <style>
            body {
                font-family: Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 600px;
                margin: 0 auto;
                padding: 20px;
            }
            .container {
                background-color: #f8f9fa;
                border-radius: 10px;
                padding: 30px;
                box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            }
            .header {
                text-align: center;
                margin-bottom: 30px;
            }
            .logo {
                font-size: 2rem;
                font-weight: bold;
                color: #007bff;
            }
            .ai-text {
                color: #2b0ae6;
                font-style: italic;
            }
            h1 {
                color: #007bff;
            }
            .credentials {
                background-color: #e9ecef;
                border-radius: 5px;
                padding: 15px;
                margin-bottom: 20px;
            }
            .btn {
                display: inline-block;
                background-color: #007bff;
                color: #ffffff;
                text-decoration: none;
                padding: 10px 20px;
                border-radius: 5px;
                margin-top: 20px;
            }
            .footer {
                text-align: center;
                margin-top: 30px;
                font-size: 0.9rem;
                color: #6c757d;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <div class="logo">Calcul<span class="ai-text">AI</span></div>
            </div>
            <h1>¡Bienvenido a CalculAI, {{ nombre_completo }}!</h1>
            <p>Estamos emocionados de que comiences tu prueba de 14 días con nosotros.</p>
            <p>Para acceder al sistema, utiliza las siguientes credenciales:</p>
            <div class="credentials">
                <p><strong>Nombre de usuario:</strong> {{ nombre_usuario }}</p>
                <p><strong>Contraseña:</strong> {{ password }}</p>
            </div>
            <p>Durante tu período de prueba, tendrás acceso completo a todas las funcionalidades del sistema.</p>
            <a href="{{ login_link }}" class="btn">Iniciar sesión en CalculAI</a>
            <p>Si tienes alguna duda o necesitas ayuda, no dudes en contactarnos.</p>
            <p>¡Gracias por elegir CalculAI!</p>
            <div class="footer">
                <p>Este es un correo automático, por favor no respondas a este mensaje.</p>
            </div>
        </div>
    </body>
    </html>
    """, nombre_completo=nombre_completo, nombre_usuario=nombre_usuario, password=password, login_link=login_link)

    data = {
        'Messages': [
            {
                "From": {
                    "Email": "soporte@sendiu.net",
                    "Name": "CalculAI"
                },
                "To": [
                    {
                        "Email": email,
                        "Name": nombre_completo
                    }
                ],
                "Subject": "Bienvenido a CalculAI - Tu período de prueba de 14 días ha comenzado",
                "HTMLPart": html_content
            }
        ]
    }

    try:
        result = mailjet.send.create(data=data)
        return result.status_code, result.json()
    except Exception as e:
        logger.error(f"Error al enviar correo de bienvenida: {str(e)}")
        return 500, {"error": str(e)}

class CustomJSONEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, datetime):
            return obj.isoformat()
        return super().default(obj)
    
def get_assistant_context():
    with current_app.app_context():
        empresa_count = Empresa.query.count()
        usuario_count = Usuario.query.count()
        return {
            "empresa_count": empresa_count,
            "usuario_count": usuario_count,
        }    

def verificar_items():
    items = InventarioItem.query.order_by(InventarioItem.id).all()
    for item in items:
        print(f"ID: {item.id}, Código: {item.codigo}, Nombre: {item.nombre}")

class AsistenteVirtual:
    def __init__(self, api_key, get_context_func):
        self.client = OpenAI(api_key=api_key)
        self.asistente_bancario = None
        self.asistente_inventario = None
        self.context = """Eres un asistente virtual multifuncional para CalculAI con acceso completo 
        a los módulos bancario e inventario. Puedes proporcionar información detallada sobre:
        - Módulo Bancario: bancos, cuentas, transacciones, balances y análisis financiero
        - Módulo Inventario: productos, stock, movimientos de inventario, almacenes, categorías y reportes"""
        
        # Inicializar los asistentes con el contexto de aplicación actual
        with current_app.app_context():
            self.inicializar_asistentes()

    def inicializar_asistentes(self):
        """Inicializa los componentes del asistente"""
        try:
            self.asistente_bancario = AsistenteBancario()
            logging.info("AsistenteBancario inicializado correctamente")
            bancos = self.asistente_bancario.obtener_info_bancos()
            logging.info(f"Bancos obtenidos: {len(bancos)}")
        except Exception as e:
            logging.error(f"Error al inicializar AsistenteBancario: {str(e)}")
            self.asistente_bancario = None

        try:
            self.asistente_inventario = AsistenteInventario()
            logging.info("AsistenteInventario inicializado correctamente")
            items = self.asistente_inventario.obtener_items()
            logging.info(f"Items obtenidos: {len(items)}")
        except Exception as e:
            logging.error(f"Error al inicializar AsistenteInventario: {str(e)}")
            self.asistente_inventario = None

    def object_as_dict(self, obj):
        def serialize(value):
            if isinstance(value, datetime):
                return value.isoformat()
            return value

        return {c.key: serialize(getattr(obj, c.key))
                for c in inspect(obj).mapper.column_attrs}

    def get_db_info(self, usuario_id):
        info = {}
        usuario = Usuario.query.get(usuario_id)
        if usuario:
            info['usuario'] = self.object_as_dict(usuario)
            info['empresa'] = self.object_as_dict(usuario.empresa) if usuario.empresa else None
        return info

    def detect_module(self, pregunta: str) -> str:
        """Detecta a qué módulo corresponde la pregunta"""
        pregunta = pregunta.lower()
        
        inventario_keywords = [
            'item', 'items', 'inventario', 'stock', 'producto', 'productos',
            'almacen', 'almacenes', 'categoria', 'categorias', 'movimiento',
            'entrada', 'salida', 'ajuste'
        ]
        
        banco_keywords = [
            'banco', 'bancos', 'cuenta', 'cuentas', 'transaccion', 'transacciones',
            'balance', 'deposito', 'transferencia', 'credito', 'debito'
        ]
        
        inventario_matches = sum(1 for keyword in inventario_keywords if keyword in pregunta)
        banco_matches = sum(1 for keyword in banco_keywords if keyword in pregunta)
        
        if inventario_matches > banco_matches:
            return 'inventario'
        elif banco_matches > inventario_matches:
            return 'banco'
        else:
            return 'general'

    def get_module_info(self, modulo: str) -> Dict[str, Any]:
        """Obtiene la información relevante según el módulo detectado"""
        try:
            if modulo == 'inventario' and self.asistente_inventario:
                items = self.asistente_inventario.obtener_items()
                items_info = []
                
                # Crear una lista formateada de items para el contexto
                for item in items:
                    item_info = {
                        "nombre": item.get("nombre", ""),
                        "codigo": item.get("codigo", ""),
                        "categoria": item.get("categoria", "Sin categoría"),
                        "stock": item.get("stock", 0),
                        "stock_minimo": item.get("stock_minimo", 0),
                        "precio": item.get("precio", 0),
                        "costo": item.get("costo", 0)
                    }
                    items_info.append(item_info)

                # Obtener el resto de la información
                resumen = self.asistente_inventario.obtener_resumen_inventario()
                almacenes = self.asistente_inventario.obtener_almacenes()
                categorias = self.asistente_inventario.obtener_categorias()
                bajo_stock = self.asistente_inventario.obtener_items_bajo_stock()

                return {
                    "items_total": len(items),
                    "items": items_info,  # Lista detallada de items
                    "resumen": resumen,
                    "almacenes": almacenes,
                    "categorias": categorias,
                    "bajo_stock": bajo_stock
                }

            elif modulo == 'banco' and self.asistente_bancario:
                bancos = self.asistente_bancario.obtener_info_bancos()
                return {
                    "bancos": bancos,
                    "bancos_total": len(bancos)
                }
            return {}
        except Exception as e:
            logging.error(f"Error al obtener información del módulo {modulo}: {str(e)}", exc_info=True)
            return {
                "error": f"Error al acceder a la información del módulo {modulo}",
                "details": str(e)
            }

    def responder(self, pregunta: str, usuario_id: int) -> Dict[str, Any]:
        try:
            logging.info(f"Procesando pregunta: {pregunta}")
            
            modulo = self.detect_module(pregunta)
            logging.info(f"Módulo detectado: {modulo}")
            
            # Verificar estado de los asistentes
            if modulo == 'banco' and not self.asistente_bancario:
                return {
                    "tipo": "texto",
                    "respuesta": "El módulo bancario no está disponible en este momento."
                }
            elif modulo == 'inventario' and not self.asistente_inventario:
                return {
                    "tipo": "texto",
                    "respuesta": "El módulo de inventario no está disponible en este momento."
                }
            
            # Obtener información del módulo
            info_modulo = self.get_module_info(modulo)
            
            # Crear contexto específico según el módulo
            if modulo == 'inventario':
                context_especifico = f"""Estás analizando el módulo de inventario.
                Información actual del inventario:
                - Total de items: {info_modulo.get('items_total', 0)}
                - Items con bajo stock: {len(info_modulo.get('bajo_stock', []))}
                - Total de almacenes: {len(info_modulo.get('almacenes', []))}
                - Total de categorías: {len(info_modulo.get('categorias', []))}
                
                Lista detallada de items en inventario:
                {json.dumps(info_modulo.get('items', []), indent=2, ensure_ascii=False)}
                
                Cuando se solicite información sobre items específicos, usa la lista anterior."""
            elif modulo == 'banco':
                context_especifico = f"""Estás analizando el módulo bancario.
                Información actual bancaria:
                - Total de bancos: {info_modulo.get('bancos_total', 0)}
                - Detalles de bancos: {json.dumps(info_modulo.get('bancos', []), ensure_ascii=False)}"""
            else:
                context_especifico = "Estás analizando una consulta general."

            context_updated = f"{self.context}\n\n{context_especifico}"
            
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": context_updated},
                    {"role": "user", "content": pregunta}
                ]
            )
            
            respuesta_texto = response.choices[0].message.content

            # Determinar si se requiere formato de lista
            palabras_clave_lista = ['lista', 'listado', 'muestra', 'enumera', 'dame', 'mostrar']
            es_solicitud_lista = any(palabra in pregunta.lower() for palabra in palabras_clave_lista)

            if es_solicitud_lista or "reporte" in pregunta.lower():
                datos = self.convertir_respuesta_a_tabla(respuesta_texto, info_modulo, modulo)
                if datos:
                    return {
                        "tipo": "lista",
                        "titulo": self.generar_titulo_lista(pregunta, modulo),
                        "contenido": respuesta_texto,
                        "datos": datos,
                        "permite_descarga": True,
                        "formatos_descarga": ["excel", "pdf", "word"]
                    }
            
            return {
                "tipo": "texto",
                "respuesta": respuesta_texto
            }

        except Exception as e:
            logging.error(f"Error al responder: {str(e)}", exc_info=True)
            return {
                "tipo": "texto",
                "respuesta": "Ha ocurrido un error al procesar tu consulta. Por favor, intenta de nuevo."
            }

    # ... (resto de los métodos sin cambios: convertir_respuesta_a_tabla, generar_titulo_lista, etc.)

    def convertir_respuesta_a_tabla(self, respuesta: str, info_modulo: Dict[str, Any], modulo: str) -> List[List[str]]:
        """Convierte una respuesta en texto a formato de tabla"""
        lineas = respuesta.split('\n')
        tabla = []
        
        if any(linea.strip().startswith(('-', '*', '•', '1.')) for linea in lineas):
            if modulo == 'inventario':
                if "stock" in respuesta.lower():
                    encabezados = ["Item", "Código", "Stock Actual", "Stock Mínimo", "Stock Máximo"]
                elif "movimiento" in respuesta.lower():
                    encabezados = ["Fecha", "Tipo", "Item", "Cantidad", "Usuario"]
                else:
                    encabezados = ["Item", "Código", "Categoría", "Stock", "Precio"]
            elif modulo == 'banco':
                if "banco" in respuesta.lower():
                    encabezados = ["Banco", "Información", "Detalles"]
                elif "cuenta" in respuesta.lower():
                    encabezados = ["Cuenta", "Balance", "Última Actualización"]
                elif "transacción" in respuesta.lower():
                    encabezados = ["Fecha", "Tipo", "Monto", "Descripción"]
                else:
                    encabezados = ["Item", "Descripción", "Valor"]
            else:
                encabezados = ["Item", "Descripción", "Valor"]
            
            tabla.append(encabezados)
            
            for linea in lineas:
                linea = linea.strip()
                if linea and not linea.startswith(('#', '=', '-', '*', '•')):
                    for marcador in ['-', '*', '•', '1.', '2.', '3.']:
                        if linea.startswith(marcador):
                            linea = linea[len(marcador):].strip()
                    
                    partes = linea.split(' - ') if ' - ' in linea else [linea]
                    while len(partes) < len(encabezados):
                        partes.append("")
                    
                    tabla.append(partes[:len(encabezados)])
            
            return tabla if len(tabla) > 1 else None
        
        return None

    def generar_titulo_lista(self, pregunta: str, modulo: str) -> str:
        """Genera un título apropiado para la lista basado en la pregunta y el módulo"""
        palabras_clave = {
            'inventario': {
                "item": "Listado de Items",
                "producto": "Listado de Productos",
                "stock": "Estado del Stock",
                "almacen": "Información de Almacenes",
                "categoria": "Listado de Categorías",
                "movimiento": "Registro de Movimientos"
            },
            'banco': {
                "banco": "Información Bancaria",
                "cuenta": "Listado de Cuentas",
                "transaccion": "Registro de Transacciones",
                "balance": "Balance de Cuentas",
                "movimiento": "Movimientos Bancarios",
            }
        }

        modulo_keywords = palabras_clave.get(modulo, {})
        for palabra, titulo in modulo_keywords.items():
            if palabra in pregunta.lower():
                return titulo

        return "Información Solicitada"

    def generar_archivos(self, usuario_id, formato):
        info = self.get_db_info(usuario_id)
        if formato == 'pdf':
            return self.generar_pdf(info)
        elif formato == 'excel':
            return self.generar_excel(info)
        elif formato == 'word':
            return self.generar_word(info)
        else:
            raise ValueError(f"Formato no soportado: {formato}")

    def generar_pdf(self, info):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        elements.append(Paragraph("Información del Usuario", styles['Title']))
        data = [["Campo", "Valor"]]
        for key, value in info['usuario'].items():
            data.append([key, str(value)])
        
        if info['empresa']:
            elements.append(Paragraph("Información de la Empresa", styles['Title']))
            for key, value in info['empresa'].items():
                data.append([key, str(value)])
        
        table = Table(data)
        table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), '#cccccc')]))
        elements.append(table)
        
        doc.build(elements)
        pdf = buffer.getvalue()
        buffer.close()
        return base64.b64encode(pdf).decode('utf-8')

    def generar_excel(self, info):
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output)
        worksheet = workbook.add_worksheet()

        row = 0
        worksheet.write(row, 0, "Información del Usuario")
        row += 1
        for key, value in info['usuario'].items():
            worksheet.write(row, 0, key)
            worksheet.write(row, 1, str(value))
            row += 1

        if info['empresa']:
            row += 1
            worksheet.write(row, 0, "Información de la Empresa")
            row += 1
            for key, value in info['empresa'].items():
                worksheet.write(row, 0, key)
                worksheet.write(row, 1, str(value))
                row += 1

        workbook.close()
        excel = output.getvalue()
        output.close()
        return base64.b64encode(excel).decode('utf-8')

    def generar_word(self, info):
        doc = Document()
        doc.add_heading('Información del Usuario', 0)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Campo'
        hdr_cells[1].text = 'Valor'
        for key, value in info['usuario'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)

        if info['empresa']:
            doc.add_heading('Información de la Empresa', 0)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Campo'
            hdr_cells[1].text = 'Valor'
            for key, value in info['empresa'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)

        buffer = BytesIO()
        doc.save(buffer)
        word = buffer.getvalue()
        buffer.close()
        return base64.b64encode(word).decode('utf-8')

def setup_logging(app):
    if not app.debug:
        file_handler = RotatingFileHandler('calculai.log', maxBytes=10240, backupCount=10)
        file_handler.setFormatter(logging.Formatter(
            '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
        ))
        file_handler.setLevel(logging.INFO)
        app.logger.addHandler(file_handler)

    app.logger.setLevel(logging.INFO)
    app.logger.info('CalculAI startup')
    return True

def create_app():
    app = Flask(__name__)
    app.config.from_object(Config)
    
    if not os.getenv('MJ_APIKEY_PUBLIC') or not os.getenv('MJ_APIKEY_PRIVATE'):
        app.logger.error('Mailjet API keys are not set. Please check your .env file.')
    
    app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'postgresql://postgres:0001@localhost:5432/calculai_db')
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
    
    from common.models import ItemFactura, ItemPreFactura, MovimientoInventario
    
    from inventario.inventario_models import InventarioItem
    from facturas.facturas_models import Facturacion, PreFactura, NotaCredito, NotaDebito, Cliente
        
    
    
    from facturas import facturacion_bp, init_facturacion
    app.register_blueprint(facturacion_bp, url_prefix='/facturacion')
    
    from marketing.routes import marketing
    app.register_blueprint(marketing, url_prefix='/marketing')
    
    from contabilidad import contabilidad_bp
    app.register_blueprint(contabilidad_bp, url_prefix='/contabilidad')
    
    from compras import compras_bp
    app.register_blueprint(compras_bp, url_prefix='/compras')

    from rrhh.routes import rrhh_bp
    app.register_blueprint(rrhh_bp, url_prefix='/rrhh')
    

    from importacion.routes import importacion_bp
    app.register_blueprint(importacion_bp)

    from proyectos.routes import proyectos_bp
    app.register_blueprint(proyectos_bp)

    from impuestos.routes import impuestos_bp
    app.register_blueprint(impuestos_bp)

    from cxc.routes import cxc_bp
    app.register_blueprint(cxc_bp)

    from cxp.routes import cxp_bp
    app.register_blueprint(cxp_bp)
        
    from activos_fijos.routes import activos_fijos_bp
    app.register_blueprint(activos_fijos_bp) 

    
    
    db.init_app(app)
    migrate.init_app(app, db)
    login_manager.init_app(app)
    login_manager.login_view = 'login'
    limiter.init_app(app)
    csrf = CSRFProtect(app)
    
    from inventario import init_app as init_inventario
    init_inventario(app)
    
    app.config['WTF_CSRF_ENABLED'] = False
    
    from admin_routes import admin
    app.register_blueprint(admin, url_prefix='/admin')
    from banco import banco_bp
    app.register_blueprint(banco_bp, url_prefix='/api')

    with app.app_context():
        from banco.banco_models import NuevoBanco as Banco
        from models import Usuario, Transaccion, Notificacion, Empresa, Rol, Permiso, Modulo, UsuarioModulo, Cuenta
        from marketing.models import Contact, Campaign, CampaignMetrics
        db.create_all()
        
    OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
    if not OPENAI_API_KEY:
        app.logger.error('OpenAI API key is not set. Please check your .env file.')    

    app.config['ASISTENTE_ACTIVO'] = True
    
    with app.app_context():
        try:
            # Ahora podemos pasar None como get_context_func ya que no lo usamos
            app.asistente = AsistenteVirtual(OPENAI_API_KEY, None)
            logger.info("Asistente virtual inicializado correctamente")
            
            # Verificar estado de inicialización
            estado = {
                "bancario": app.asistente.asistente_bancario is not None,
                "inventario": app.asistente.asistente_inventario is not None
            }
            logger.info(f"Estado de inicialización de asistentes: {estado}")
            
        except Exception as e:
            logger.error(f"Error al inicializar el asistente virtual: {str(e)}")
            app.asistente = None

    # Configurar el logging
    setup_logging(app)

    @app.errorhandler(Exception)
    def handle_exception(e):
        app.logger.error(f"Unhandled exception: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500
    
    asistente_notificacion = signal('asistente-notificacion')

    @app.route('/admin/toggle_asistente_usuario/<int:user_id>', methods=['POST'])
    @login_required
    def toggle_asistente_usuario(user_id):
        if current_user.rol != 'admin':
            return jsonify({"error": "No tienes permisos para realizar esta acción"}), 403
        
        usuario = Usuario.query.get_or_404(user_id)
        usuario.asistente_activo = not usuario.asistente_activo
        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": f"Asistente {'activado' if usuario.asistente_activo else 'desactivado'} para {usuario.nombre_usuario}",
            "asistente_activo": usuario.asistente_activo
        })
     
    @app.route('/api/asistente/verificar-inventario')
    @login_required
    def verificar_inventario():
        if not app.asistente:
            return jsonify({"error": "Asistente no inicializado"})
        
        resultado = app.asistente.verificar_datos_inventario()
        return jsonify(resultado) 
        
    @app.route('/api/asistente/diagnostico')
    @login_required
    def diagnostico_asistente():
        if not app.asistente:
            return jsonify({
                "status": "error",
                "message": "Asistente no inicializado"
            })
        
        estado = {
            "asistente_bancario": {
                "inicializado": app.asistente.asistente_bancario is not None,
                "bancos_count": len(app.asistente.asistente_bancario.obtener_info_bancos()) if app.asistente.asistente_bancario else 0
            },
            "asistente_inventario": {
                "inicializado": app.asistente.asistente_inventario is not None,
                "items_count": len(app.asistente.asistente_inventario.obtener_items()) if app.asistente.asistente_inventario else 0
            }
        }
        
        return jsonify(estado)   

    @asistente_notificacion.connect
    def log_notificacion(sender, mensaje):
        print(f"Notificación al asistente: {mensaje}")
    
    @event.listens_for(Banco, 'after_insert')
    def notificar_nuevo_banco(mapper, connection, target):
        asistente_notificacion.send('Banco', mensaje=f"Nuevo banco creado: {target.nombre}")

    @event.listens_for(Banco, 'after_update')
    def notificar_actualizacion_banco(mapper, connection, target):
        asistente_notificacion.send('Banco', mensaje=f"Banco actualizado: {target.nombre}") 
        
    # Agregar estos manejadores de eventos después de los existentes de banco
    @event.listens_for(InventarioItem, 'after_insert')
    def notificar_nuevo_item(mapper, connection, target):
        asistente_notificacion.send('Inventario', mensaje=f"Nuevo item creado: {target.nombre}")

    @event.listens_for(InventarioItem, 'after_update')
    def notificar_actualizacion_item(mapper, connection, target):
        asistente_notificacion.send('Inventario', mensaje=f"Item actualizado: {target.nombre}")

    @event.listens_for(MovimientoInventario, 'after_insert')
    def notificar_movimiento_inventario(mapper, connection, target):
        tipo_mov = "entrada" if target.tipo == "entrada" else "salida"
        asistente_notificacion.send('Inventario', 
            mensaje=f"Nuevo movimiento de {tipo_mov} registrado para item ID: {target.item_id}")       
    
    @app.before_request
    def log_request_info():
        app.logger.debug('Headers: %s', request.headers)
        app.logger.debug('Body: %s', request.get_data())
    
    @login_manager.user_loader
    def load_user(user_id):
        return Usuario.query.get(int(user_id))

    @app.route('/')
    def index():
        return render_template('index.html')

    @app.route('/login', methods=['GET', 'POST'])
    @limiter.limit("50 per minute")
    def login():
        if current_user.is_authenticated:
            return redirect(url_for('index'))

        if request.method == 'POST':
            username = request.form.get('username')
            password = request.form.get('password')

            logger.debug(f"Received login attempt - Username: {username}, Password: {'*' * len(password) if password else None}")

            if not username or not password:
                logger.warning("Intento de inicio de sesión con credenciales incompletas")
                return jsonify({
                    "success": False,
                    "error": "Por favor, proporcione un nombre de usuario y contraseña"
                }), 400

            user = Usuario.query.filter_by(nombre_usuario=username).first()

            if user and user.check_password(password):
                login_user(user)
                logger.info(f"Usuario {username} ha iniciado sesión exitosamente")
                return jsonify({
                    "success": True,
                    "message": "Inicio de sesión exitoso"
                }), 200
            else:
                logger.warning(f"Intento de inicio de sesión fallido para el usuario {username}")
                return jsonify({
                    "success": False,
                    "error": "Nombre de usuario o contraseña incorrectos"
                }), 401

        return render_template('login.html')

    @app.route('/api/registro_exitoso')
    def api_registro_exitoso():
        data = {
            "title": "¡Registro Exitoso!",
            "description": "Tu cuenta ha sido creada correctamente.",
            "message": "Hemos enviado un correo electrónico de confirmación.",
            "emailMessage": "Revisa tu bandeja de entrada o la carpeta de spam para confirmar tu correo.",
            "spamMessage": "Si no recibes el correo, intenta registrarte nuevamente.",
            "loginButtonText": "Iniciar Sesión"
        }
        return jsonify(data)

    @app.route('/registro_exitoso')
    def registro_exitoso():
        return render_template('registro_exitoso.html')

    @app.route("/facturacion/facturas")
    @login_required
    def facturas():
        return render_template('facturacion/facturas.html')

    @app.route('/facturacion/pre-facturas')
    @login_required
    def pre_facturas():
        return render_template('facturacion/pre_facturas.html')

    @app.route("/facturacion/reporte-de-ventas")
    @login_required
    def reporte_ventas():
        return render_template('facturacion/reporte_de_ventas.html')

    @app.route("/facturacion/gestion-de-clientes")
    @login_required
    def gestion_clientes():
        return render_template('facturacion/gestion_de_clientes.html')
    
    

    @app.route('/inventario/entrada-almacen')
    @login_required
    def inventario_entrada_almacen():
        return render_template('inventario/entrada_almacen.html')

    @app.route('/inventario/salida-almacen')
    @login_required
    def inventario_salida_almacen():
        return render_template('inventario/salida_almacen.html')

    @app.route('/inventario/inventario')
    @login_required
    def inventario_actual():
        return render_template('inventario/inventario.html')

    @app.route('/inventario/reporte')
    @login_required
    def inventario_reporte():
        return render_template('inventario/reporte.html')

    @app.route("/reset_password", methods=["POST"])
    def reset_password():
        username = request.form.get("username")
        new_password = request.form.get("new_password")
        user = Usuario.query.filter_by(nombre_usuario=username).first()

        if not user:
            logger.warning(f"Intento de restablecimiento de contraseña para usuario no existente: {username}")
            return jsonify({"success": False, "error": "Usuario no encontrado"}), 404

        if not new_password or len(new_password) < 8:
            return jsonify({"success": False, "error": "La nueva contraseña debe tener al menos 8 caracteres"}), 400
        try:
            user.set_password(new_password)
            db.session.commit()
            logger.info(f"Contraseña actualizada para el usuario {username}")
            return jsonify({"success": True, "message": "Contraseña actualizada correctamente"}), 200
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error al actualizar la contraseña para {username}: {str(e)}")
            return jsonify({"success": False, "error": "Error al actualizar la contraseña"}), 500

    @app.route("/request_password_reset", methods=["POST"])
    def request_password_reset():
        email = request.form.get("email")
        user = Usuario.query.filter_by(email=email).first()

        if not user:
            logger.warning(f"Intento de restablecimiento de contraseña para email no registrado: {email}")
            return jsonify({"success": False, "error": "No se encontró un usuario con ese correo electrónico"}), 404

        logger.info(f"Solicitud de restablecimiento de contraseña para el usuario: {user.nombre_usuario}")
        return jsonify({"success": True, "message": "Se ha enviado un correo electrónico con instrucciones para restablecer tu contraseña"}), 200

    @app.route('/logout', methods=['POST'])
    @login_required
    def logout():
        logger.info(f"Usuario {current_user.nombre_usuario} ha cerrado sesión")
        logout_user()
        return redirect(url_for('login'))
    
    def generate_password(length=12):
        alphabet = string.ascii_letters + string.digits + string.punctuation
        password = ''.join(secrets.choice(alphabet) for i in range(length))
        return password

    @app.route('/registro', methods=['GET', 'POST'])
    def registro():
        if current_user.is_authenticated:
            return redirect(url_for('index'))

        if request.method == 'POST':
            data = request.form
            logger.info(f"Datos de registro recibidos: {data}")

            required_fields = ['nombre', 'apellido', 'telefono', 'nombre_usuario', 'email', 
                            'nombre_empresa', 'rnc_empresa', 'direccion_empresa', 'telefono_empresa']
            
            missing_fields = [field for field in required_fields if not data.get(field)]
            if missing_fields:
                logger.warning(f"Campos faltantes en el registro: {', '.join(missing_fields)}")
                return jsonify({"success": False, "error": f"Faltan los siguientes campos: {', '.join(missing_fields)}"}), 400

            nombre = data.get('nombre')
            apellido = data.get('apellido')
            telefono = data.get('telefono')
            nombre_usuario = data.get('nombre_usuario')
            email = data.get('email')
            password = data.get('password')

            nombre_empresa = data.get('nombre_empresa')
            rnc_empresa = data.get('rnc_empresa')
            direccion_empresa = data.get('direccion_empresa')
            telefono_empresa = data.get('telefono_empresa')

            try:
                if Usuario.query.filter_by(nombre_usuario=nombre_usuario).first():
                    logger.warning(f"Intento de registro con nombre de usuario existente: {nombre_usuario}")
                    return jsonify({"success": False, "error": "El nombre de usuario ya está en uso"}), 400

                if Usuario.query.filter_by(email=email).first():
                    logger.warning(f"Intento de registro con email existente: {email}")
                    return jsonify({"success": False, "error": "El correo electrónico ya está registrado"}), 400

                nueva_empresa = Empresa(
                    nombre=nombre_empresa,
                    rnc=rnc_empresa,
                    direccion=direccion_empresa,
                    telefono=telefono_empresa,
                    estado="activo"
                )
                db.session.add(nueva_empresa)
                db.session.flush()

                new_user = Usuario(
                    nombre_usuario=nombre_usuario,
                    email=email,
                    nombre=nombre,
                    apellido=apellido,
                    telefono=telefono,
                    empresa_id=nueva_empresa.id,
                    asistente_activo=False
                )
                new_user.set_password(password)
                db.session.add(new_user)
                db.session.commit()

                logger.info(f"Nuevo usuario registrado: {nombre_usuario}")

                login_link = url_for('login', _external=True)
                status_code, response = send_welcome_email(email, nombre, apellido, login_link, nombre_usuario, password)

                if status_code == 200:
                    logger.info(f"Correo de bienvenida enviado a: {email}")
                    session['registro_exitoso'] = True
                    return jsonify({"success": True, "message": "Registro exitoso", "redirect": url_for('registro_exitoso')})
                else:
                    logger.error(f"Error al enviar correo de bienvenida: {response}")
                    return jsonify({"success": True, "message": "Registro exitoso, pero hubo un problema al enviar el correo de bienvenida", "redirect": url_for('registro_exitoso')})

            except IntegrityError as e:
                db.session.rollback()
                logger.error(f"Error de integridad en el registro: {str(e)}")
                return jsonify({"success": False, "error": "Error de integridad en la base de datos"}), 500
            except Exception as e:
                db.session.rollback()
                logger.error(f"Error inesperado en el registro: {str(e)}")
                return jsonify({"success": False, "error": "Ocurrió un error inesperado durante el registro"}), 500

        return render_template('registro.html')

    @app.route("/api/modulos")
    @login_required
    def get_modulos():
        modulos = [
            "Banco",
            "Contabilidad",
            "Activos Fijos",
            "Cuentas Por Cobrar",
            "Cuentas Por Pagar",
            "Facturacion",
            "Impuestos",
            "Inventario",
            "Compras",
            "Importacion",
            "Proyectos",
            "Recursos Humanos",
            "Marketing",
        ]
        return jsonify(modulos)

    @app.route("/api/submodulos/<modulo>")
    @login_required
    def get_submodulos(modulo):
        submodulos = {
            "Contabilidad": [
                "Cuentas", "Diario", "Mayor General", "Balanza de Comprobación",
                "Estado de Resultados", "Balance General", "Configuraciones", "Flujo de caja",
            ],
            "Banco": [
                "Bancos", "Depósitos", "Notas de Crédito/Débito", "Transferencias Bancarias",
                "Conciliación Bancaria", "Gestión de Bancos", "Divisas",
            ],
            "Activos Fijos": [
                "Activo Fijo", "Depreciación", "Retiro", "Revalorización", "Tipo de Activo Fijo",
            ],
            "Cuentas Por Cobrar": [
                "Cliente", "Descuento y devoluciones", "Nota de credito", "Nota de debito",
                "Recibo", "Anticipo CxC", "Condicion de pago", "Reporte CxC", "Tipo de cliente",
            ],
            "Cuentas Por Pagar": [
                "Factura Suplidor", "Nota de Crédito", "Nota de Débito", "Orden de Compras",
                "Suplidor", "Anticipo CxP", "Pago de Contado", "Reporte CxP",
                "Requisición Cotización", "Solicitud Compras", "Tipo de Suplidor",
            ],
            "Facturacion": [
                "Facturas", "Pre-facturas",
                "Reporte de Ventas", "Gestión de clientes",
            ],
            "Impuestos": [
                "Formulario 606", "Formulario 607", "Reporte IT1",
                "Impuesto sobre la Renta (IR17)", "Serie Fiscal", "Configuraciones",
            ],
            "Inventario": [
                "Items", "Entrada de Almacén", "Salida de Almacén",
                "Inventario", "Reporte de Inventario",
            ],
            "Compras": [
                "Solicitudes de Compra", "Órdenes de Compra", "Recepción de Materiales",
                "Gastos", "Reporte de Compras/Gastos",
            ],
            "Importacion": [
                "Expediente de Importacion", "Importador", "Reportes Importacion",
            ],
            "Proyectos": [
                "Gestión de Proyectos", "Presupuestos", "Facturación por Proyecto",
            ],
            "Recursos Humanos": [
                "Gestión de Empleados", "Nómina", "Evaluación de Desempeño",
            ],
            "Marketing": [
                "Gestión de Contactos", "Campañas de Email", "Plantillas de Email",
                "Reportes de Campañas", "Segmentación de Contactos",
                "Automatizaciones", "Integración de Redes Sociales",
            ],
        }
        return jsonify(submodulos.get(modulo, []))

    @app.route("/transacciones")
    @login_required
    def transacciones():
        transacciones = Transaccion.obtener_todos()
        return render_template("transacciones.html", transacciones=transacciones)

    @app.route("/api/buscar-transaccion")
    @login_required
    def buscar_transaccion():
        tipo = request.args.get("tipo")
        descripcion = request.args.get("descripcion")
        cuenta_id = request.args.get("cuenta_id")

        transacciones = Transaccion.buscar(tipo=tipo, descripcion=descripcion, cuenta_id=cuenta_id)
        logger.info(f"Búsqueda de transacciones realizada. Resultados: {len(transacciones)}")
        return jsonify([transaccion.to_dict() for transaccion in transacciones])

    @app.route("/api/crear-transaccion", methods=["POST"])
    @login_required
    def crear_transaccion():
        datos = request.json
        try:
            nueva_transaccion = Transaccion(
                tipo=datos["tipo"],
                monto=datos["monto"],
                descripcion=datos["descripcion"],
                cuenta_id=datos["cuenta_id"],
            )
            db.session.add(nueva_transaccion)
            db.session.commit()
            logger.info(f"Nueva transacción creada: {nueva_transaccion.id}")
            return jsonify(nueva_transaccion.to_dict()), 201
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error al crear transacción: {str(e)}")
            return jsonify({"error": f"Error al crear la transacción: {str(e)}"}), 500

    @app.route("/api/tareas")
    @login_required
    def get_tareas():
        tareas = [
            {"descripcion": "Revisar facturas pendientes", "vence": "2023-08-15"},
            {"descripcion": "Preparar informe mensual", "vence": "2023-08-20"},
            {"descripcion": "Reunión con inversores", "vence": "2023-08-25"},
        ]
        return jsonify(tareas)

    @app.route("/api/notificaciones")
    @login_required
    def get_notificaciones():
        notificaciones = [
            {"mensaje": "Nuevo cliente registrado", "tipo": "info"},
            {"mensaje": "Factura #1234 vencida", "tipo": "warning"},
            {"mensaje": "Actualización del sistema disponible", "tipo": "info"},
        ]
        return jsonify(notificaciones)

    @app.route('/api/asistente', methods=['POST'])
    @login_required
    def consultar_asistente():
        try:
            if not app.config['ASISTENTE_ACTIVO'] or not current_user.asistente_activo:
                return jsonify({
                    "tipo": "texto",
                    "respuesta": "El asistente no está activo para tu usuario. Contacta al administrador."
                }), 403

            pregunta = request.json.get('pregunta', '').strip()
            if not pregunta:
                return jsonify({
                    "tipo": "texto",
                    "respuesta": "Por favor, proporciona una pregunta."
                }), 400
            
            # Obtener la respuesta del asistente
            respuesta = app.asistente.responder(pregunta, current_user.id)
            
            # Asegurarnos de que la respuesta tenga el formato correcto
            if isinstance(respuesta, dict):
                if respuesta.get('tipo') == 'lista':
                    return jsonify({
                        "tipo": "lista",
                        "titulo": respuesta.get('titulo', 'Resultados'),
                        "contenido": respuesta.get('contenido', ''),
                        "datos": respuesta.get('datos', []),
                        "permite_descarga": respuesta.get('permite_descarga', True),
                        "formatos_descarga": respuesta.get('formatos_descarga', ["excel", "pdf", "word"])
                    })
                elif respuesta.get('tipo') == 'texto':
                    return jsonify({
                        "tipo": "texto",
                        "respuesta": respuesta.get('respuesta', '')
                    })
            
            # Si la respuesta no tiene el formato esperado
            return jsonify({
                "tipo": "texto",
                "respuesta": str(respuesta)
            })

        except Exception as e:
            logger.error(f"Error en consultar_asistente: {str(e)}")
            return jsonify({
                "tipo": "texto",
                "respuesta": "Ha ocurrido un error inesperado. Por favor, intenta de nuevo más tarde."
            }), 500

    @app.route('/api/asistente_status')
    @login_required
    def asistente_status():
        return jsonify({
            "activo": current_user.asistente_activo and app.config['ASISTENTE_ACTIVO']
        })

    @app.route('/api/actualizar_estado_asistente', methods=['POST'])
    @login_required
    def actualizar_estado_asistente():
        if current_user.rol != 'admin':
            return jsonify({"error": "No tienes permisos para realizar esta acción"}), 403
        
        nuevo_estado = request.json.get('activo', False)
        app.config['ASISTENTE_ACTIVO'] = nuevo_estado
        return jsonify({"mensaje": "Estado del asistente actualizado correctamente"})
    
    @app.route("/api/datos_graficos")
    @login_required
    def get_datos_graficos():
        datos = {
            "ventas": {
                "labels": ["Ene", "Feb", "Mar", "Abr", "May", "Jun"],
                "values": [100, 200, 150, 300, 250, 400],
            },
            "ingresos_vs_gastos": {
                "labels": ["Ene", "Feb", "Mar", "Abr", "May", "Jun"],
                "ingresos": [1000, 1200, 1100, 1300, 1250, 1400],
                "gastos": [800, 900, 850, 950, 900, 1000],
            },
            "distribucion": {
                "labels": ["Ventas", "Gastos", "Beneficios"],
                "values": [400, 300, 100],
            },
        }
        return jsonify(datos)

    @app.route("/api/usuario")
    @login_required
    def get_usuario():
        usuario = {
            "nombre": f"{current_user.nombre} {current_user.apellido}",
            "id": current_user.id,
            "avatar": url_for('placeholder_image', width=100, height=100),
            "empresa": current_user.empresa.nombre if current_user.empresa else "Sin empresa"
        }
        return jsonify(usuario)

    @app.route("/api/placeholder/<int:width>/<int:height>")
    def placeholder_image(width, height):
        return (
            f"Placeholder image of {width}x{height}",
            200,
            {"Content-Type": "text/plain"},
        )

    @app.route('/favicon.ico')
    def favicon():
        try:
            return send_from_directory(os.path.join(app.root_path, 'static'),
                                    'favicon.ico', mimetype='image/vnd.microsoft.icon')
        except FileNotFoundError:
            app.logger.warning("Favicon not found. Returning 404.")
            return '', 404

    @app.route("/admin_panel")
    @login_required
    def admin_panel():
        if current_user.rol != "admin":
            flash("Acceso no autorizado", "error")
            return redirect(url_for("index"))
        return render_template("admin_panel.html")

    @app.route("/Bancos")
    @login_required
    def redirect_to_bancos():
        return redirect(url_for('banco.sub_bancos'))
    
    @app.route("/depositos")
    @login_required
    def redirect_to_depositos():
        return render_template('banco/depositos.html')
    
    @app.route("/notas-credito-debito")
    @login_required
    def notas_credito_debito():
        return render_template('banco/notas_credito_debito.html')

    @app.route("/transferencias-bancarias")
    @login_required
    def transferencias_bancarias():
        return render_template('banco/transferencias.html')

    @app.route("/conciliacion-bancaria")
    @login_required
    def conciliacion_bancaria():
        return render_template('banco/conciliacion.html')

    @app.route("/gestion-bancos")
    @login_required
    def gestion_bancos():
        return render_template('banco/gestion_bancos.html')

    @app.route("/divisas")
    @login_required
    def divisas():
        return render_template('banco/divisas.html')

    @app.route('/api/generar-contenido-email', methods=['POST'])
    @login_required
    def generar_contenido_email():
        data = request.json
        prompt = data.get('prompt')
        
        if not prompt:
            return jsonify({"error": "Se requiere un prompt"}), 400 
        
        try:
            ai_prompt = f"""Actúa como un experto en marketing digital y diseño de emails. 
            Crea una plantilla de email HTML profesional y atractiva basada en la siguiente descripción: {prompt}.La plantilla debe ser responsive y utilizar estilos en línea para compatibilidad con clientes de email.
            Incluye contenido relevante y convincente que se ajuste al tema solicitado.
            Para las imágenes, utiliza placeholders de https://placehold.co con dimensiones y texto descriptivo apropiados. 
            Asegúrate de que el diseño sea moderno, atractivo y optimizado para conversiones.
            IMPORTANTE: Devuelve ÚNICAMENTE el código HTML de la plantilla, sin ningún texto explicativo adicional."""
            
            contenido_generado = app.asistente.responder(ai_prompt, current_user.id)
            
            html_content = extraer_html(contenido_generado['contenido'])
            
            return jsonify({"contenido": html_content})
        except Exception as e:
            logger.error(f"Error al generar contenido de email con IA: {str(e)}")
            return jsonify({"error": "Ocurrió un error al generar el contenido"}), 500

    def extraer_html(contenido):
        import re
        match = re.search(r'<html>[\s\S]*?</html>', contenido)
        if match:
            return match.group(0)
        else:
            return contenido

    @app.route('/api/descargar_reporte/<tipo_archivo>', methods=['POST'])
    @login_required
    def descargar_reporte(tipo_archivo):
        datos = request.json.get('datos')
        if not datos:
            return jsonify({"error": "No se proporcionaron datos para el reporte"}), 400

        if tipo_archivo == 'csv':
            return generar_csv(datos)
        elif tipo_archivo == 'pdf':
            return generar_pdf(datos)
        elif tipo_archivo == 'word':
            return generar_word(datos)
        else:
            return jsonify({"error": "Formato de archivo no soportado"}), 400

    def generar_csv(datos):
        output = StringIO()
        writer = csv.writer(output)
        writer.writerows(datos)
        output.seek(0)
        return send_file(
            BytesIO(output.getvalue().encode('utf-8')),
            mimetype='text/csv',
            as_attachment=True,
            attachment_filename='reporte.csv'
        )

    def generar_pdf(datos):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        
        t = Table(datos)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 12),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(t)
        
        doc.build(elements)
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            attachment_filename='reporte.pdf'
        )

    def generar_word(datos):
        document = Document()
        document.add_heading('Reporte', 0)

        table = document.add_table(rows=1, cols=len(datos[0]))
        hdr_cells = table.rows[0].cells
        for i, item in enumerate(datos[0]):
            hdr_cells[i].text = str(item)

        for row in datos[1:]:
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            attachment_filename='reporte.docx'
        )

    @app.errorhandler(CSRFError)
    def handle_csrf_error(e):
        return jsonify({"error": "CSRF token missing or incorrect"}), 400

    @app.errorhandler(404)
    def page_not_found(e):
        return render_template('404.html'), 404

    @app.errorhandler(500)
    def internal_server_error(e):
        return render_template('500.html'), 500

    @app.after_request
    def add_header(response):
        response.headers['X-Content-Type-Options'] = 'nosniff'
        response.headers['X-Frame-Options'] = 'SAMEORIGIN'
        response.headers['X-XSS-Protection'] = '1; mode=block'
        return response
 


    return app

if __name__ == '__main__':
    app = create_app()
    app.run(debug=True)